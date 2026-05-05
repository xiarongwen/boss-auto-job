#!/usr/bin/env python3
"""
简历解析模块 - 支持 PDF 和 Word 格式

Usage:
  from parse_resume import parse_resume
  text = parse_resume("/path/to/resume.pdf")
  text = parse_resume("/path/to/resume.docx")
"""

import sys
from pathlib import Path


def parse_resume(file_path: str) -> str:
    """解析简历文件，返回纯文本内容。

    支持格式:
      - PDF (.pdf)
      - Word (.docx, .doc)
      - 纯文本 (.txt)

    Args:
        file_path: 简历文件路径

    Returns:
        简历的纯文本内容

    Raises:
        FileNotFoundError: 文件不存在
        ValueError: 不支持的文件格式
        ImportError: 缺少必要的解析库
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"简历文件不存在: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(path)
    elif suffix == ".docx":
        return _parse_docx(path)
    elif suffix == ".doc":
        return _parse_doc(path)
    elif suffix == ".txt":
        return path.read_text(encoding="utf-8").strip()
    else:
        raise ValueError(f"不支持的文件格式: {suffix}（支持 PDF/DOCX/DOC/TXT）")


def _parse_pdf(path: Path) -> str:
    """解析 PDF 简历。

    优先使用 pymupdf (速度快、质量好)，降级到 pdfplumber。
    """
    # 方法1: pymupdf (fitz)
    try:
        import fitz  # pymupdf
        doc = fitz.open(str(path))
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        result = "\n".join(text_parts).strip()
        if result:
            return result
    except ImportError:
        pass
    except Exception as e:
        print(f"[warn] pymupdf 解析失败: {e}", file=sys.stderr)

    # 方法2: pdfplumber
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        result = "\n".join(text_parts).strip()
        if result:
            return result
    except ImportError:
        pass
    except Exception as e:
        print(f"[warn] pdfplumber 解析失败: {e}", file=sys.stderr)

    # 方法3: PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(path))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        result = "\n".join(text_parts).strip()
        if result:
            return result
    except ImportError:
        pass
    except Exception as e:
        print(f"[warn] PyPDF2 解析失败: {e}", file=sys.stderr)

    raise ImportError(
        "无法解析 PDF，请安装以下任一库:\n"
        "  pip install pymupdf\n"
        "  pip install pdfplumber\n"
        "  pip install PyPDF2"
    )


def _parse_docx(path: Path) -> str:
    """解析 Word (.docx) 简历。"""
    # 方法1: python-docx
    try:
        from docx import Document
        doc = Document(str(path))
        text_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)
        # 也提取表格中的内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        result = "\n".join(text_parts).strip()
        if result:
            return result
    except ImportError:
        pass
    except Exception as e:
        print(f"[warn] python-docx 解析失败: {e}", file=sys.stderr)

    raise ImportError(
        "无法解析 DOCX，请安装:\n"
        "  pip install python-docx"
    )


def _parse_doc(path: Path) -> str:
    """解析旧版 Word (.doc) 简历。

    .doc 是二进制格式，需要先转换。
    """
    import subprocess
    import tempfile
    import shutil

    # 方法1: 用 LibreOffice 转换
    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    if libreoffice:
        with tempfile.TemporaryDirectory() as tmpdir:
            try:
                subprocess.run(
                    [libreoffice, "--headless", "--convert-to", "docx",
                     "--outdir", tmpdir, str(path)],
                    capture_output=True, timeout=30, check=True
                )
                converted = Path(tmpdir) / (path.stem + ".docx")
                if converted.exists():
                    return _parse_docx(converted)
            except Exception as e:
                print(f"[warn] LibreOffice 转换失败: {e}", file=sys.stderr)

    # 方法2: textract (需要系统依赖)
    try:
        import textract
        text = textract.process(str(path)).decode("utf-8")
        if text.strip():
            return text.strip()
    except ImportError:
        pass
    except Exception as e:
        print(f"[warn] textract 解析失败: {e}", file=sys.stderr)

    # 方法3: antiword
    try:
        result = subprocess.run(
            ["antiword", str(path)],
            capture_output=True, text=True, timeout=10
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except FileNotFoundError:
        pass
    except Exception:
        pass

    raise ImportError(
        "无法解析 .doc 格式，请尝试以下方案:\n"
        "  1. 将 .doc 另存为 .docx 格式\n"
        "  2. 安装 LibreOffice: brew install --cask libreoffice\n"
        "  3. 安装 antiword: brew install antiword"
    )


def resume_to_text(resume_path: str, output_path: str = None) -> str:
    """解析简历并可选保存为文本文件。

    Args:
        resume_path: 简历文件路径
        output_path: 可选，保存解析结果的文本文件路径

    Returns:
        简历文本内容
    """
    text = parse_resume(resume_path)

    if output_path:
        Path(output_path).write_text(text, encoding="utf-8")
        print(f"[ok] 简历文本已保存: {output_path}")

    return text


# CLI 入口
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="解析简历文件 (PDF/Word) 为纯文本")
    parser.add_argument("file", help="简历文件路径 (.pdf/.docx/.doc/.txt)")
    parser.add_argument("-o", "--output", help="保存文本到文件")
    args = parser.parse_args()

    try:
        text = resume_to_text(args.file, args.output)
        print(text)
    except (FileNotFoundError, ValueError, ImportError) as e:
        print(f"[error] {e}", file=sys.stderr)
        sys.exit(1)
